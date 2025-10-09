# Criação de arquivos word do zero!

from docx import Document
from docx.shared import Cm

doc = Document()
doc.add_heading('Título do documento',0)

para = doc.add_paragraph('Um parágrafo simples')
para.add_run(' e importante').bold = True
para.add_run(' com palavras em')
para.add_run('italico').italic = True

doc.add_heading('Titulo - Nível 1', level = 1)
doc.add_heading('Titulo - Nível 2', level = 2)
doc.add_heading('Titulo - Nível 3', level = 3)

doc.add_paragraph('Formatação "No Spacing"',style='No Spacing')
doc.add_paragraph('Formatação "Heading 1"',style='Heading 1')
doc.add_paragraph('Formatação "Heading 2"',style='Heading 2')
doc.add_paragraph('Formatação "Heading 3"',style='Heading 3')
doc.add_paragraph('Formatação "Title"',style='Title')
doc.add_paragraph('Formatação "Subtitle"',style='Subtitle')
doc.add_paragraph('Formatação "Intense Quote"',style='Intense Quote')
doc.add_paragraph('Formatação "List Paragraph"',style='List Paragraph')
doc.add_paragraph('Formatação "List Bullet"',style='List Bullet')
doc.add_paragraph('Formatação "List Number"',style='List Number')

doc.add_picture('d:\\Estudos\\TI\\Programação\\Linguagens\\Python\\curso\\modulo9\\documentos_word\\notebook.jpg', width=Cm(5.25))

tabela = doc.add_table(rows=3, cols=2)
celula1 = tabela.cell(0,0)
celula1.text = 'Nome'

compras = (
    (3, '101', 'Uva'),
    (7, '422', 'Pão'),
    (4, '423', 'Banana, Ovos, Tomate , Care')
)

# Criar a estrutura inicial da tabela 
tabela_supermercado = doc.add_table(rows=1, cols=3)
cabecalho_table_supermercado = tabela_supermercado.rows[0].cells
cabecalho_table_supermercado[0].text = 'Quantidade'
cabecalho_table_supermercado[1].text = 'Id'
cabecalho_table_supermercado[2].text = 'Descrição'

for quantidade,id,desc in compras:
    linha = tabela_supermercado.add_row().cells
    linha[0].text = str(quantidade)
    linha[1].text = id
    linha[2].text = desc

doc.add_page_break()
# Para salver em um novo arquivo eu tenho que mudar o nome do arquivo 
doc.save('d:\\Estudos\\TI\\Programação\\Linguagens\\Python\\curso\\modulo9\\documentos_word\\demo.docx')

